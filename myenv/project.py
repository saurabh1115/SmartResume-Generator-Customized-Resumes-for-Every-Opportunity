import streamlit as st
import openai
import os
import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv
import google.generativeai as genai

# Load environment variables
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")  # Ensure this is properly set in your .env file

# Check if API key is available
if not GOOGLE_API_KEY:
    st.error("Google API key not found. Please check your .env file.")
else:
    genai.configure(api_key=GOOGLE_API_KEY)

# Streamlit UI setup
st.title("SmartResume Generator: Customized Resumes for Every Opportunity")
st.write("Fill in your details to generate a professional and customized resume.")

# User Inputs
name = st.text_input("Full Name", "")
email = st.text_input("Email", "")
phone = st.text_input("Phone", "")
linkedin = st.text_input("LinkedIn URL", "")
summary = st.text_area("Professional Summary", "")

# Experience Section
st.subheader("Work Experience")
experience = []
num_experiences = st.number_input("Number of experiences", min_value=1, max_value=10, step=1)

for i in range(num_experiences):
    st.markdown(f"### Experience {i+1}")
    job_title = st.text_input(f"Job Title {i+1}", key=f"job_title_{i}")
    company = st.text_input(f"Company {i+1}", key=f"company_{i}")
    duration = st.text_input(f"Duration {i+1}", key=f"duration_{i}")
    description = st.text_area(f"Description {i+1}", key=f"description_{i}")
    experience.append({
        "job_title": job_title,
        "company": company,
        "duration": duration,
        "description": description
    })

# Skills Section
skills = st.text_area("Skills (comma-separated)", "")

# Education Section
st.subheader("Education")
degree = st.text_input("Degree", "")
university = st.text_input("University", "")
grad_year = st.text_input("Graduation Year", "")

# Additional Details
st.subheader("Additional Details (Optional)")
certifications = st.text_area("Certifications", "")
languages = st.text_area("Languages Spoken", "")
projects = st.text_area("Key Projects", "")

# Generate Resume Function
def generate_resume():
    try:
        model = genai.GenerativeModel("gemini-pro")
        prompt = f"""
        Generate a professional resume with proper formatting:
        Name: {name}
        Email: {email}
        Phone: {phone}
        LinkedIn: {linkedin}
        Summary: {summary}
        Experience: {experience}
        Skills: {skills}
        Education: {degree}, {university}, {grad_year}
        Certifications: {certifications}
        Languages Spoken: {languages}
        Key Projects: {projects}
        Format the response properly.
        """
        
        response = model.generate_content(prompt)
        return response.text if response else "Error generating resume."

    except Exception as e:
        return f"An error occurred: {str(e)}"

# Save Resume to Word Document
def save_to_word(resume_text):
    try:
        doc = docx.Document()
        doc.add_heading(name, level=1)

        # Contact Information
        contact_info = f"{email} | {phone} | {linkedin}"
        doc.add_paragraph(contact_info).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Summary
        doc.add_heading("Professional Summary", level=2)
        doc.add_paragraph(summary)

        # Experience
        doc.add_heading("Work Experience", level=2)
        for exp in experience:
            doc.add_paragraph(f"{exp['job_title']} at {exp['company']} ({exp['duration']})", style="Heading 3")
            doc.add_paragraph(exp['description'])

        # Skills
        doc.add_heading("Skills", level=2)
        doc.add_paragraph(skills)

        # Education
        doc.add_heading("Education", level=2)
        doc.add_paragraph(f"{degree}, {university} ({grad_year})")

        # Certifications
        if certifications:
            doc.add_heading("Certifications", level=2)
            doc.add_paragraph(certifications)

        # Languages
        if languages:
            doc.add_heading("Languages Spoken", level=2)
            doc.add_paragraph(languages)

        # Key Projects
        if projects:
            doc.add_heading("Key Projects", level=2)
            doc.add_paragraph(projects)

        # Save file
        file_path = "Generated_Resume.docx"
        doc.save(file_path)
        return file_path
    except Exception as e:
        return f"Error saving resume: {str(e)}"

# Generate Resume Button
if st.button("Generate Resume"):
    resume_text = generate_resume()
    st.text_area("Generated Resume", resume_text, height=400)

    file_path = save_to_word(resume_text)
    if file_path:
        with open(file_path, "rb") as f:
            st.download_button("Download Resume", f, file_name="Generated_Resume.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
